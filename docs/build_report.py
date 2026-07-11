"""
Builds docs/Nautilus-ERP-System-Report.docx from facts verified against this repository.

Run:  python docs/build_report.py
Requires: pip install python-docx

Every claim in the generated document was checked against source in src/ and client/.
Nothing here is aspirational: where a capability does not exist, the document says so.

The document is written in two parts. Part I is for any reader — an owner, a manager,
an auditor — and avoids assuming technical background. Part II is a technical reference
for developers and IT staff.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
LOGO = REPO / "client" / "public" / "nautilus-logo.png"
OUT = REPO / "docs" / "Nautilus-ERP-System-Report.docx"


def logo_stream():
    """
    A fresh stream per use: python-docx consumes the file object, and the logo is placed
    twice (title page and page header).
    """
    from io import BytesIO

    return BytesIO(LOGO.read_bytes())

TEAL = RGBColor(0x0E, 0x73, 0x67)
INK = RGBColor(0x14, 0x31, 0x2B)
MUTED = RGBColor(0x5A, 0x6B, 0x66)

TABLE_STYLE = "Light Grid Accent 1"


# --------------------------------------------------------------------------- helpers
def field(paragraph, instruction: str):
    """Insert a Word field code (used for TOC and page numbers)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ""
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def para(doc, text, *, italic=False, size=None, color=None, align=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = TABLE_STYLE
    except KeyError:
        t.style = "Table Grid"
    t.autofit = True
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(head)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)


def part_page(doc, kicker, title, subtitle):
    """A visual divider page between Part I and Part II."""
    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph()
    para(doc, kicker, size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = TEAL
    para(doc, subtitle, size=11, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 18), ("Heading 2", 13.5)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = TEAL if name == "Heading 1" else INK
        st.paragraph_format.space_before = Pt(16 if name == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(6)


def add_header_logo(section):
    """Small logo, right-aligned, in the running header."""
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = p.add_run("Nautilus ERP    ")
    label.font.size = Pt(8)
    label.font.color.rgb = MUTED
    p.add_run().add_picture(logo_stream(), width=Inches(0.28))


def add_footer_page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    field(p, "PAGE")
    run2 = p.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.color.rgb = MUTED
    field(p, "NUMPAGES")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED


# --------------------------------------------------------------------------- document
def build():
    doc = Document()
    style_document(doc)

    first = doc.sections[0]
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(first, margin, Inches(1))

    # ---------------- title page (no header/footer chrome)
    doc.add_paragraph()
    doc.add_paragraph()
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(logo_stream(), width=Inches(2.0))

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("Nautilus ERP")
    tr.font.size = Pt(34)
    tr.bold = True
    tr.font.color.rgb = TEAL

    para(
        doc,
        "Business Management System — System Report",
        size=14,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    para(
        doc,
        "An integrated sales, purchasing and inventory platform for a Fiji-registered business",
        size=11,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
    doc.add_paragraph()
    para(
        doc,
        "Purpose: to describe, accurately and without exaggeration, what this system does today, "
        "who may use each part of it, and what it does not yet do. Part I is written for any "
        "reader; Part II is a technical reference for developers and IT staff.",
        size=10.5,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    para(doc, "11 July 2026", size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        doc,
        "Prepared from the source code of the repository, not from a specification.",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )

    # ---------------- body section with header/footer
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(body, margin, Inches(1))
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False
    add_header_logo(body)
    add_footer_page_numbers(body)
    # keep the title page clean
    first.footer.is_linked_to_previous = False
    first.header.is_linked_to_previous = False

    # ---------------- how to read this document
    h1(doc, "How to read this document")
    table(
        doc,
        ["If you are…", "Read…", "You can safely skip…"],
        [
            (
                "An owner or director deciding whether to rely on this system",
                "Sections 1, 3, 7 and 8 — especially 8.1, the known limitations",
                "Part II entirely",
            ),
            (
                "A manager or staff member who will use it daily",
                "Sections 2 to 6, and the glossary in section 9 when a term is unfamiliar",
                "Part II entirely",
            ),
            (
                "An auditor or accountant",
                "Sections 3, 4, 7 and 8; section 13 (data and audit trail) if you want the mechanics",
                "Sections 5, 6, 15",
            ),
            (
                "A developer or IT administrator",
                "All of Part II (sections 10–16); Part I for context",
                "Sections 5 and 6",
            ),
        ],
        widths=[2.1, 2.9, 1.9],
    )
    caption(doc, "Table 1 — A reading guide by audience.")

    # ---------------- table of contents
    h1(doc, "Contents")
    toc_p = doc.add_paragraph()
    field(toc_p, r'TOC \o "1-3" \h \z \u')
    caption(
        doc,
        "This table of contents is a live Word field. If it appears empty or out of date, "
        "click anywhere inside it and press F9 (or Ctrl+A then F9) to refresh it.",
    )

    part_page(
        doc,
        "PART I",
        "The system, for everyone",
        "What Nautilus ERP does, who may use it, and what it does not do. No technical background assumed.",
    )

    # ============================================================ 1
    h1(doc, "1. Executive summary")
    para(
        doc,
        "Nautilus ERP is an enterprise resource planning system — a single application in which a "
        "business records what it sells, what it buys, and what it holds in stock, instead of "
        "spreading those records across spreadsheets and separate tools. It is built for a "
        "Fiji-registered trading business: a supermarket, wholesaler, distributor or pharmacy that "
        "issues VAT tax invoices in Fiji dollars and needs an auditable record of every stock "
        "movement and every payment.",
    )
    para(
        doc,
        "The system covers six areas of the business. A catalogue holds products, categories, units "
        "of measure, currencies and tax rates. An inventory module tracks stock in warehouses and "
        "values it using FIFO — first-in, first-out, meaning the oldest stock you bought is treated "
        "as the first stock you sold. A sales module carries a customer order through to an issued "
        "invoice and the payments against it. A purchasing module does the mirror image: a purchase "
        "order, the goods receipt when the delivery arrives, the supplier's bill and the payment. A "
        "reporting module summarises the position and exports it. An administration module manages "
        "users, roles, branches and company details, and exposes an audit trail of every change.",
    )
    para(
        doc,
        "Three independent layers of access control sit over all of this. Roles decide what kind of "
        "action a person may take. Branch scoping decides which records they can see at all. "
        "Segregation of duties — five maker-checker rules — prevents one person from completing a "
        "purchase-to-payment chain on their own, which is the classic route by which a business is "
        "defrauded by its own staff. Customers never sign in; this is a staff system.",
    )
    para(
        doc,
        "Two limitations deserve to appear this early, because a report that buries them is not worth "
        "reading. First, the system does not submit invoices to the Fiji Revenue and Customs Service. "
        "The integration point for FRCS fiscalization exists and is honest about itself: every issued "
        "invoice is recorded as NotSubmitted. Nothing is faked as accredited. Second, Nautilus ERP is "
        "not an accounting system. It has no general ledger and no double-entry bookkeeping. It "
        "records the operational truth of the business and can feed an accountant; it does not replace "
        "one. Section 8.1 sets out the full list of known limitations.",
    )

    # ============================================================ 2
    h1(doc, "2. What the system does")
    para(
        doc,
        "The following six modules make up the working system. Each is implemented end to end: a "
        "domain model, application logic, an HTTP endpoint, and a screen in the browser client.",
    )

    table(
        doc,
        ["Module", "What it holds", "What you can do with it"],
        [
            (
                "Catalogue and reference data",
                "Products, categories, units of measure, currencies, taxes and tax rates, branches, warehouses",
                "Maintain the shared vocabulary every other module depends on; set VAT rates with effective dates",
            ),
            (
                "Inventory",
                "Stock levels per product and warehouse, FIFO cost layers, an immutable stock-movement ledger",
                "Receive, issue, transfer and adjust stock; set reorder levels; view stock valuation and full movement history",
            ),
            (
                "Sales",
                "Customers, sales orders, invoices with line-level VAT, customer payments",
                "Take an order, fulfil it from stock, raise and issue a tax invoice, record part or full payment",
            ),
            (
                "Purchasing",
                "Suppliers, purchase orders, goods receipts, supplier invoices, supplier payments",
                "Raise and approve a purchase order, receive goods into stock, approve the supplier's bill, pay it",
            ),
            (
                "Reporting",
                "Dashboard key figures, inventory valuation, sales trend",
                "Read the current position; export the inventory valuation report as CSV, Excel or PDF; print a Fiji tax invoice as PDF",
            ),
            (
                "Administration",
                "User accounts, roles, branch assignment, company profile, audit log",
                "Create users, assign roles and a branch, activate or deactivate accounts, edit company and tax details, review who changed what",
            ),
        ],
        widths=[1.6, 2.2, 2.7],
    )
    caption(doc, "Table 2 — The six modules of Nautilus ERP.")
    para(
        doc,
        "A seventh capability sits underneath these six rather than beside them: any record in any "
        "module can carry file attachments — a scanned supplier invoice, a delivery photo, a signed "
        "contract — uploaded, listed, downloaded and removed by referencing the record they belong to, "
        "without a schema change per module. The screen is live today on the four document detail pages "
        "— sales orders, customer invoices, purchase orders, supplier invoices — as an Attachments card "
        "alongside each document's lines; other record types (products, customers, suppliers) can use "
        "the same API but do not yet have the screen wired in. Section 11 documents the endpoints.",
    )

    h2(doc, "2.1 Catalogue and reference data")
    para(
        doc,
        "Nothing else works until the catalogue is right. A product carries its category, its unit of "
        "measure, its pricing and the tax that applies to it. Tax itself is modelled properly rather "
        "than as a number on an invoice: a tax has a treatment — Standard, ZeroRated or Exempt — and a "
        "Standard tax carries a history of effective-dated rates. Changing Fiji's VAT rate is therefore "
        "a data operation performed by an administrator, not a change to the software.",
    )

    h2(doc, "2.2 Inventory")
    para(
        doc,
        "Stock is held per product per warehouse. Every change to a quantity on hand is written as an "
        "entry in a ledger that is never edited or deleted: a receipt, an issue, a positive or negative "
        "stock-take adjustment, or the two halves of a transfer between warehouses. Because the ledger "
        "is immutable, the stock position at any past date can be reconstructed and defended.",
    )
    para(
        doc,
        "Valuation uses FIFO cost layers. When stock arrives, a layer is created recording the quantity "
        "and the unit cost paid. When stock goes out, the oldest layers are consumed first and the cost "
        "of goods sold is taken from them. A transfer between warehouses preserves the original cost, "
        "so moving stock does not invent or destroy value. Attempting to issue, transfer or adjust more "
        "stock than exists is refused outright; nothing is half-posted.",
    )

    h2(doc, "2.3 Sales, purchasing and reporting")
    para(
        doc,
        "Sales and purchasing are described in detail in section 4, because their value lies in the "
        "sequence of states each document passes through, and section 6 walks through both chains as a "
        "story. Reporting is deliberately narrow: a dashboard "
        "of key figures (customer, supplier and product counts, inventory value, low-stock count, sales "
        "this month, money owed by customers and to suppliers, open orders on both sides), a sales trend, "
        "and an inventory valuation report exportable as CSV, Excel or PDF. A customer invoice can also "
        "be rendered as a Fiji-format PDF tax invoice.",
    )

    # ============================================================ 3
    h1(doc, "3. Who uses it, and what they can see")
    para(
        doc,
        "Access is decided by three separate mechanisms, applied one after the other. A request must "
        "pass all three. They answer three different questions: what kind of action is this person "
        "allowed to perform; which records are they allowed to touch; and does this particular action "
        "conflict with something they already did on the same document.",
    )

    h2(doc, "3.1 Roles")
    para(
        doc,
        "There are three roles. Every signed-in user holds at least one. Customers and suppliers do not "
        "have accounts and cannot sign in — Nautilus ERP is an internal staff system, and there is no "
        "public self-registration of any kind.",
    )
    table(
        doc,
        ["Role", "Can read", "Can change", "Cannot"],
        [
            (
                "Administrator",
                "Everything, including the audit trail",
                "Everything: all transactions, plus users, roles, branch assignment, company profile and reference data",
                "Bypass segregation-of-duties rules; those apply regardless of role",
            ),
            (
                "Manager",
                "All operational data within their branch scope",
                "Transactions and reference data: stock movements, sales orders, invoices, purchase orders, goods receipts, supplier invoices and payments",
                "Manage user accounts or roles; view the audit trail; edit the company profile",
            ),
            (
                "Staff",
                "Operational data within their branch scope — orders, invoices, stock levels, reports",
                "One transactional action: posting a goods receipt when a delivery arrives, since the person at the loading dock is rarely a manager. Otherwise read-only in the transactional modules; may update their own profile and password",
                "Post other stock movements, confirm or approve documents, record payments, manage users",
            ),
        ],
        widths=[1.1, 1.9, 2.4, 1.6],
    )
    caption(
        doc,
        "Table 3 — The three roles. Write access to sales, purchasing and inventory endpoints is "
        "restricted to Administrator and Manager, with one deliberate exception: Staff may post goods "
        "receipts, and segregation of duties still bars them from receiving an order they raised or approved.",
    )

    h2(doc, "3.2 Branch scoping")
    para(
        doc,
        "A role says what you may do; it does not say which records you may do it to. That is branch "
        "scoping, and it works at the level of the individual record. When a user is assigned to a "
        "branch, their sign-in token carries a branch claim. Warehouse-bound data — stock, movements, "
        "sales orders, purchase orders — belongs to a branch through its warehouse, so a branch-scoped "
        "user sees and acts on only the warehouses of their branch. A user with no branch assigned is "
        "unrestricted, which is the correct default for a head-office administrator. A user whose branch "
        "has no warehouses sees nothing, which is the correct behaviour for a misconfigured account: it "
        "fails closed.",
    )

    h2(doc, "3.3 Segregation of duties")
    para(
        doc,
        "Segregation of duties is the principle that no single person should be able to complete a "
        "transaction chain alone. The fraud it prevents is well known: raise a purchase order to a "
        "supplier you control, approve your own order, record goods that never arrived, approve the "
        "invoice for them, and pay it. Nautilus ERP breaks that chain in five places. Each rule compares "
        "the person attempting the current step against the people who performed the earlier, conflicting "
        "steps on the same document chain.",
    )
    table(
        doc,
        ["Rule", "The person acting may not be…", "What it stops"],
        [
            (
                "PurchaseOrderApproval",
                "the person who raised the purchase order",
                "Approving your own commitment to spend money",
            ),
            (
                "GoodsReceipt",
                "the person who raised or approved the order",
                "Confirming the arrival of goods you ordered and signed off",
            ),
            (
                "SupplierInvoiceApproval",
                "the person who entered the bill, or who received the goods",
                "Certifying a liability you created or vouched for",
            ),
            (
                "SupplierPayment",
                "the person who approved the supplier's bill",
                "Releasing money against a bill you yourself approved",
            ),
            (
                "InvoiceVoid",
                "the person who issued the customer invoice",
                "Quietly cancelling a sale you recorded — a common way to conceal cash theft",
            ),
        ],
        widths=[1.9, 2.6, 2.5],
    )
    caption(doc, "Table 4 — The five segregation-of-duties rules, enforced by default.")
    para(
        doc,
        "A blocked attempt returns a 403 Forbidden response with an explanation, and is written to the "
        "log — a refused attempt is itself information worth keeping. The rules are on by default. A "
        "business too small to staff both sides of a control can switch individual rules off, or all of "
        "them, through configuration, accepting the audit trail as the compensating control. Documents "
        "created with no signed-in user, such as those from the demo data seeder, have no recorded actor "
        "and therefore never trip a rule.",
    )

    # ============================================================ 4
    h1(doc, "4. The document lifecycles")
    para(
        doc,
        "A document in an ERP is not a form; it is a small state machine. Each document may only move "
        "between defined states, and an attempt at an illegal move is refused with a 409 Conflict rather "
        "than being quietly allowed. This is what makes the records trustworthy after the fact.",
    )

    h2(doc, "4.1 Order to cash")
    para(
        doc,
        "A customer places an order. The order is captured as a Draft, which is still editable. Confirming "
        "it commits the business to the sale. Fulfilling it issues the stock — an all-or-nothing operation "
        "that consumes FIFO cost layers and writes the ledger entries; if there is not enough stock, "
        "nothing is posted and the order simply remains Confirmed. An invoice is then raised from the "
        "order, snapshotting the VAT rate in force on that date so an issued invoice never changes "
        "retroactively. Issuing the invoice makes it a committed tax document. Payments received against "
        "it move it to PartiallyPaid and then to Paid. An invoice can be voided only before any payment "
        "has been recorded, and never by the person who issued it.",
    )

    h2(doc, "4.2 Procure to pay")
    para(
        doc,
        "The purchasing chain is the mirror image. A purchase order starts as a Draft, is confirmed and "
        "placed with the supplier, and then receives goods — possibly across several deliveries, moving "
        "it to PartiallyReceived and finally Received. Posting a goods receipt is the point at which "
        "stock actually enters the warehouse and a FIFO cost layer is created at the purchase-order unit "
        "cost. Over-receiving a line is refused and nothing is posted. The supplier's bill is then raised "
        "from the order, capturing input VAT, approved, and paid. Four of the five segregation-of-duties "
        "rules police this chain.",
    )

    table(
        doc,
        ["Document", "States", "Notes on the transitions"],
        [
            (
                "Sales order",
                "Draft → Confirmed → Fulfilled; Cancelled",
                "Fulfilment issues stock via FIFO; all-or-nothing. Cancelled is terminal.",
            ),
            (
                "Customer invoice",
                "Draft → Issued → PartiallyPaid → Paid; Void",
                "Issuing snapshots the VAT rate and calls the fiscalization port. Void only before any payment.",
            ),
            (
                "Purchase order",
                "Draft → Confirmed → PartiallyReceived → Received; Cancelled",
                "Goods receipts drive the received states. Cancellation only before receipt.",
            ),
            (
                "Supplier invoice",
                "Draft → Approved → PartiallyPaid → Paid; Cancelled",
                "Approved is a committed liability. Cancellation only before payment.",
            ),
            (
                "Stock movement",
                "Receipt, Issue, AdjustmentIn, AdjustmentOut, TransferIn, TransferOut",
                "Not a lifecycle but a ledger: entries are immutable and never revised.",
            ),
        ],
        widths=[1.4, 2.4, 3.1],
    )
    caption(doc, "Table 5 — Document state machines. Illegal transitions are rejected with 409 Conflict.")

    # ============================================================ 5
    h1(doc, "5. A tour of the screens")
    para(
        doc,
        "The browser client is organised around a sidebar of modules. What a given person sees depends "
        "on their role — an administration screen simply does not appear for a Staff account — but the "
        "full set of screens is as follows.",
    )
    table(
        doc,
        ["Screen", "Who sees it", "What it is for"],
        [
            (
                "Sign in / Forgot password / Reset password",
                "Everyone",
                "Email-and-password sign-in, with a code-entry step when the account has two-factor authentication on; a reset flow that never reveals whether an account exists",
            ),
            (
                "Dashboard",
                "All roles",
                "The key figures at a glance: inventory value, low-stock alerts, sales this month, receivables and payables, open orders, and a sales trend chart",
            ),
            (
                "Products",
                "All roles (read); Admin and Manager edit",
                "The product catalogue: SKU, category, unit of measure, prices, applicable tax, reorder level",
            ),
            (
                "Inventory",
                "All roles (read); Admin and Manager post",
                "Stock on hand per warehouse, valuation, and the movement ledger; post receipts, issues, transfers and stock-take adjustments",
            ),
            (
                "Customers / Suppliers",
                "All roles (read); Admin and Manager edit",
                "The trading partners: contact details, TIN, and each partner's document history",
            ),
            (
                "Sales orders → order detail",
                "All roles (read); Admin and Manager act",
                "Capture an order, confirm it, fulfil it from stock, raise the invoice from it; attach files (any signed-in role)",
            ),
            (
                "Invoices → invoice detail",
                "All roles (read); Admin and Manager act",
                "Issue the tax invoice, record payments against it, download the Fiji-format PDF, void (subject to the rules in section 3.3); attach files (any signed-in role)",
            ),
            (
                "Purchase orders → order detail",
                "All roles (read); Staff may post receipts; Admin and Manager act",
                "Raise and confirm an order, record goods receipts as deliveries arrive; attach files (any signed-in role)",
            ),
            (
                "Supplier invoices → detail",
                "All roles (read); Admin and Manager act",
                "Enter the supplier's bill against the order, approve it, record payments; attach files (any signed-in role)",
            ),
            (
                "Reports",
                "All roles",
                "The inventory valuation report with CSV, Excel and PDF export",
            ),
            (
                "Profile",
                "The signed-in user",
                "Update your own name, contact details and password; set up or turn off two-factor authentication",
            ),
            (
                "Users",
                "Administrator only",
                "Create accounts, assign roles and a branch, activate or deactivate",
            ),
            (
                "Company / Taxes / Reference data / Settings",
                "Administrator only",
                "The company profile and TIN; taxes and effective-dated rates; categories, units, currencies, branches and warehouses",
            ),
            (
                "Audit log",
                "Administrator only",
                "Every insert, update and delete of a business record: who, when, and what changed",
            ),
        ],
        widths=[1.8, 1.7, 3.4],
    )
    caption(doc, "Table 6 — The screens of the browser client, by audience.")
    para(
        doc,
        "Real-time notifications appear as they happen — an invoice issued, stock received — delivered "
        "over a live connection rather than by refreshing the page. One honest caveat, expanded in "
        "section 8.1: today those notifications are broadcast to every signed-in staff member rather "
        "than targeted at the people they concern.",
    )

    # ============================================================ 6
    h1(doc, "6. Two walkthroughs")
    para(
        doc,
        "The mechanics above are easier to hold onto as a story. The names are invented; every step, "
        "permission and refusal described is real system behaviour.",
    )

    h2(doc, "6.1 Buying stock: procure to pay, with four people")
    numbered(
        doc,
        [
            "Mere, a Manager at the Suva branch, notices on the Dashboard that flour is below its "
            "reorder level. She raises a purchase order on Pacific Flour Mills for 200 bags and confirms "
            "it. Because she raised it, the system will not let her approve it — that is the "
            "PurchaseOrderApproval rule.",
            "Jone, another Manager, reviews and approves the order. It is now a commitment to spend.",
            "The truck arrives on Thursday. Sela, a Staff member on the loading dock, counts 120 bags "
            "and posts a goods receipt for exactly that quantity — the one transactional action Staff "
            "may perform. Stock enters the warehouse at the order's unit cost, a FIFO layer is created, "
            "and the order becomes PartiallyReceived. Had Sela been the one who raised or approved the "
            "order, the GoodsReceipt rule would have refused her. The remaining 80 bags arrive the "
            "following week and a second receipt moves the order to Received.",
            "The supplier's bill arrives by email. Mere enters it against the order, capturing the input "
            "VAT. Jone approves it — Sela could not have, having received the goods, and Mere could not "
            "approve a bill she entered (SupplierInvoiceApproval).",
            "Finance records the payment. Jone, who approved the bill, is barred from paying it "
            "(SupplierPayment); Mere records the payment instead. Four people, no one of whom could "
            "have run the chain alone — and every step, including the refusals, is in the audit trail.",
        ],
    )

    h2(doc, "6.2 Selling stock: order to cash")
    numbered(
        doc,
        [
            "A hotel phones in a weekly order. Mere captures it as a Draft sales order — still "
            "editable — then confirms it, committing the stock promise.",
            "She fulfils the order. The system consumes the oldest FIFO cost layers for each line and "
            "writes immutable ledger entries; if any line were short of stock, nothing at all would "
            "post and the order would stay Confirmed.",
            "She raises the invoice from the order. Each line snapshots the VAT rate in force that day, "
            "so a future rate change cannot silently alter an old invoice. She issues it — it is now a "
            "tax document, printed as a Fiji-format PDF with the company TIN and its fiscal status "
            "(NotSubmitted; see section 7.1) on its face.",
            "The hotel pays half on delivery and the balance at month end: the invoice moves to "
            "PartiallyPaid and then Paid. Had the sale gone wrong before any payment, the invoice could "
            "be voided — but not by Mere, who issued it (InvoiceVoid).",
        ],
    )

    # ============================================================ 7
    h1(doc, "7. Fiji localization")
    para(
        doc,
        "The system assumes a Fiji-registered business as its default case rather than treating Fiji as a "
        "configuration afterthought. Three things follow from that.",
    )
    para(
        doc,
        "VAT. Fiji's standard Value Added Tax rate is 15 percent. It is stored as a percentage — the "
        "number 15, meaning 15%, not the fraction 0.15 — and it is effective-dated. A tax carries a "
        "history of rate rows, each with a start date and an optional end date; the rate in force on a "
        "given date is the one whose window contains it. Adding a new rate closes the current open one. "
        "Zero-rated and exempt supplies are modelled as distinct treatments, not merely as a rate of "
        "zero, because the distinction matters for revenue reporting. Every invoice line snapshots the "
        "rate that applied on the day of issue.",
    )
    para(
        doc,
        "The tax invoice. A company profile holds the legal and trading name, the FRCS Taxpayer "
        "Identification Number (TIN), address, contact details and the base currency, which defaults to "
        "FJD. An issued invoice can be rendered as a PDF laid out as a Fiji tax invoice: seller name and "
        "TIN, buyer name and TIN, line items with their snapshotted VAT, totals in Fiji dollars, and the "
        "invoice's fiscalization status printed plainly on the face of the document. Payment methods "
        "include the tender types actually used in Fiji, mobile wallet among them.",
    )

    h2(doc, "7.1 FRCS / VMS fiscalization — the honest status")
    para(
        doc,
        "This system does not submit invoices to the Fiji Revenue and Customs Service, and it does not "
        "pretend to.",
    )
    para(
        doc,
        "Fiscalization — the electronic accreditation of invoices with the revenue authority through "
        "FRCS's TPOS / VAT Monitoring System — is present in the architecture as a port: a named boundary "
        "with a defined contract, where a real adapter will one day be plugged in. What ships today is a "
        "null stub. When an invoice is issued, the stub is called; it writes a log line saying no verified "
        "adapter is configured, and returns the status NotSubmitted. Every issued invoice therefore carries "
        "a fiscal status of NotSubmitted, and that status is shown in the interface and printed on the PDF.",
    )
    para(
        doc,
        "This is a deliberate decision, not an omission awaiting apology. The FRCS specification, including "
        "the accredited invoice-numbering format, has not been verified against an authoritative source. "
        "Writing code that reports invoices as successfully submitted when they were not would produce a "
        "system that lies to its operator about their tax compliance. The boundary is modelled so that "
        "substituting a verified adapter is a single dependency-injection change and requires no rework of "
        "the sales module. Until then, a business using Nautilus ERP must meet its FRCS fiscalization "
        "obligations by other means.",
    )

    # ============================================================ 8
    h1(doc, "8. Security posture")
    para(
        doc,
        "Signing in issues two tokens. A short-lived access token, valid for fifteen minutes, accompanies "
        "each request. A longer-lived refresh token exchanges for a new access token when that expires. "
        "Refresh tokens rotate on every use: the old one is consumed and replaced, and presenting an "
        "already-consumed token is treated as evidence of theft and rejected. Refresh tokens are stored "
        "hashed with SHA-256 rather than in the clear, so a reader of the database cannot impersonate a "
        "user. Passwords must be at least eight characters with upper case, lower case, a digit and a "
        "symbol; five failed attempts lock the account for fifteen minutes; every attempt, successful or "
        "not, is written to a login history. Password reset and forgotten-password responses never reveal "
        "whether an account exists.",
    )
    para(
        doc,
        "Every account may add a second factor: a standard TOTP authenticator app (Google Authenticator, "
        "Authy and similar). Turning it on requires proving possession of the authenticator with a live "
        "code before it takes effect, so a mistaken setup can never lock the account out, and it issues a "
        "set of one-time recovery codes for the case where the device is lost. Once enabled, a correct "
        "password no longer signs a user in on its own: it returns a short-lived challenge, valid for five "
        "minutes and scoped so it cannot be used as a bearer token anywhere else, which must then be "
        "redeemed with a current authenticator code or an unused recovery code before session tokens are "
        "issued. Multi-factor authentication is opt-in per account; there is no policy that requires it, "
        "for Administrator accounts or otherwise.",
    )
    para(
        doc,
        "Beyond authentication, the API applies rate limiting (tighter on the authentication endpoints than "
        "elsewhere), response compression, baseline security headers, HTTPS redirection outside development, "
        "and a startup guard that refuses to boot outside development if the JWT signing key is missing, "
        "too short, or still the development default. All database access is through Entity Framework with "
        "parameterized queries; there is no string-concatenated SQL. Every insert, update and delete of a "
        "business entity is captured by an audit interceptor and readable, by administrators only, as an "
        "audit trail. Deletion is soft: records are marked deleted, not removed.",
    )

    h2(doc, "8.1 Known limitations and what is not implemented")
    para(
        doc,
        "The following are true of the system as it stands. They are stated plainly because a reader who "
        "discovers them later would be right to distrust everything else in this document.",
    )
    bullets(
        doc,
        [
            "Multi-factor authentication is available but not required of anyone, including Administrator accounts — there is no policy that enforces it.",
            "No email verification of accounts, and no public self-registration — accounts are created by an administrator, which mitigates that but does not remove it.",
            "The refresh token is held in the browser's localStorage. This is convenient and survives a page reload, but it is readable by any script that manages to run on the page, so a cross-site scripting flaw would be more damaging than it would be with an HttpOnly cookie.",
            "File attachments have a screen on the four document detail pages (sales orders, invoices, purchase orders, supplier invoices) but not on other record types — products, customers and suppliers can be attached to over the API, but nothing in the interface calls it there yet.",
            "No general ledger and no double-entry accounting. Nautilus ERP records operational transactions — stock, invoices, payments — not journals, chart of accounts or trial balance. It is not a substitute for accounting software.",
            "FRCS / VMS fiscalization is not integrated. Every issued invoice remains NotSubmitted. See section 7.1.",
            "Real-time notifications sent through SignalR are broadcast to every connected staff member. Per-user targeting exists in the code, but the business events that fire notifications use the broadcast path, so a notification about one branch's invoice reaches everyone signed in.",
            "Emails (password resets, notifications) are queued for background delivery and sent over real SMTP when a mail server is configured; left unconfigured, delivery falls back to a stub that writes the email to the log instead, which is what a fresh checkout does until someone supplies SMTP settings.",
            "The Hangfire background job store is in-memory, so queued work does not survive a restart of the API.",
            "File attachments are written to local disk by default. That is durable on a single persistent instance but not on an ephemeral or multi-instance host; a cloud storage adapter would need to be substituted for those deployments, and the abstraction (IFileStorage) is written so that substitution needs no change to the callers.",
            "Multi-currency is modelled — FJD is the base currency and currencies are reference data — but there is no exchange-rate table in use and no bank reconciliation, RTGS/ACH or mobile-wallet payment integration. Mobile wallet is a payment method you may record, not a payment rail the system connects to.",
            "There is no offline mode. GUID primary keys keep that option open architecturally; nothing implements it.",
        ],
    )

    # ============================================================ 9
    h1(doc, "9. Glossary")
    para(
        doc,
        "Terms used in this document and in the application, in plain language.",
    )
    table(
        doc,
        ["Term", "Meaning"],
        [
            ("ERP", "Enterprise resource planning: one system of record for sales, purchasing and stock, instead of separate spreadsheets."),
            ("VAT", "Value Added Tax, Fiji's consumption tax. The standard rate is 15%. Charged on sales (output VAT) and paid on purchases (input VAT)."),
            ("TIN", "Taxpayer Identification Number, issued by FRCS. Appears on every tax invoice for both seller and buyer."),
            ("FRCS", "Fiji Revenue and Customs Service, the tax authority."),
            ("Fiscalization / VMS / TPOS", "FRCS's scheme for electronically accrediting invoices with the tax authority. Not yet integrated — see section 7.1."),
            ("FIFO", "First-in, first-out: the oldest stock purchased is treated as the first stock sold, which decides the cost recorded against each sale."),
            ("Cost layer", "A record of one batch of stock received: its quantity and the unit cost paid. FIFO consumes the oldest layers first."),
            ("Goods receipt", "The document recording that a delivery physically arrived: what was counted, into which warehouse, against which purchase order."),
            ("Sales order vs. invoice", "The order is the promise; the invoice is the tax document. Stock moves when the order is fulfilled; money is owed when the invoice is issued."),
            ("Segregation of duties", "The rule that no single person may complete a money-moving chain alone — also called maker-checker. See section 3.3."),
            ("Branch scoping", "Limiting what a user can see and touch to the warehouses of their assigned branch."),
            ("Audit trail", "An administrator-readable record of every change to business data: who, when, and what changed."),
            ("Soft delete", "Deleted records are marked as deleted and hidden, not destroyed — so history and the audit trail stay intact."),
            ("Void", "Cancelling an issued invoice before any payment. The invoice remains on record, marked Void; it does not vanish."),
            ("Draft / Confirmed / Issued …", "Document states. A document can only move along its defined path (section 4); illegal jumps are refused."),
        ],
        widths=[1.8, 5.1],
    )
    caption(doc, "Table 7 — Glossary.")

    part_page(
        doc,
        "PART II",
        "Technical reference",
        "Architecture, API surface, domain model, data, testing and operations. Written for developers and IT staff.",
    )

    # ============================================================ 10
    h1(doc, "10. Architecture and technology")
    para(
        doc,
        "Nautilus ERP is a .NET 9 solution laid out as Clean "
        "Architecture, a structure in which dependencies point inward toward the business rules. The "
        "Domain project — the entities and the rules that govern them — knows nothing of databases, web "
        "frameworks or user interfaces, and a test in the suite fails the build if anyone makes it "
        "otherwise. Around it sits an Application layer holding the use cases, then Infrastructure and "
        "Persistence for the outside world, then a thin web API whose controllers only translate HTTP "
        "into application requests and back.",
    )
    table(
        doc,
        ["Project", "Responsibility"],
        [
            ("ERP.Domain", "Entities, enums, domain rules and domain exceptions. No external dependencies."),
            ("ERP.Application", "Use cases as CQRS commands and queries, validation, authorization rules, ports (interfaces) to the outside world."),
            ("ERP.Infrastructure", "Adapters: JWT token service, SMTP/logging-stub email, fiscalization stub, SignalR notifications, Hangfire jobs, report exporters, local file storage, Sentry error tracking."),
            ("ERP.Persistence", "EF Core DbContext, entity configurations, audit interceptor, seeding, SQL Server / SQLite migrations."),
            ("ERP.Persistence.Migrations.Postgres", "The PostgreSQL migration set, kept as its own project so provider-specific migrations never mix."),
            ("ERP.API", "ASP.NET Core web API: 21 controllers, middleware, rate limiting, health checks, Swagger in development."),
            ("ERP.Shared", "Cross-cutting primitives shared by the layers (Result, Error, role names)."),
            ("client/", "React + TypeScript single-page application, built with Vite; TanStack Query for server state; Tailwind-based design system; route-level code-splitting; Vitest + React Testing Library."),
        ],
        widths=[2.2, 4.7],
    )
    caption(doc, "Table 8 — Solution layout.")
    para(
        doc,
        "Application logic follows CQRS — Command Query Responsibility Segregation, meaning operations that "
        "change data and operations that read it are separate objects — dispatched through the MediatR "
        "library, with FluentValidation running as a pipeline step before any handler executes. Expected "
        "failures are not exceptions. A handler returns a Result carrying either a value or an Error with a "
        "code, and the API maps those codes to HTTP status: validation to 400, unauthorized to 401, not "
        "found to 404, forbidden to 403, conflict to 409, locked-out account to 423. Exceptions are reserved "
        "for the genuinely unexpected.",
    )
    para(
        doc,
        "Reports are exported through a provider-agnostic exporter: CSV natively, Excel via ClosedXML, "
        "PDF via QuestPDF. Real-time notifications go over SignalR. Background work is queued through "
        "Hangfire (in-memory storage; see section 8.1).",
    )

    # ============================================================ 11
    h1(doc, "11. The API surface")
    para(
        doc,
        "The API exposes 83 endpoints across 21 controllers, all under /api and all requiring "
        "authentication except sign-in, the MFA challenge, token refresh and the password-reset pair. "
        "Interactive documentation is served at /swagger in development. The table groups the controllers "
        "by module; docs/API.md in the repository documents each endpoint individually.",
    )
    table(
        doc,
        ["Area", "Controllers (endpoints)", "Notes"],
        [
            (
                "Authentication & identity",
                "Auth (12), Users (6)",
                "Sign-in, MFA challenge and self-service setup/enable/disable, refresh, logout, password change/forgot/reset, profile; user administration is Administrator-only",
            ),
            (
                "Catalogue & reference",
                "Products (5), Categories (2), UnitsOfMeasure (2), Currencies (2), Taxes (3), Branches (2), Warehouses (2)",
                "Reads for all roles; writes for Administrator and Manager; taxes include effective-dated rate management",
            ),
            (
                "Inventory",
                "Inventory (7)",
                "Stock levels, valuation, movement history; receipt, issue, transfer and adjustment postings",
            ),
            (
                "Sales",
                "Customers (2), SalesOrders (6), Invoices (8)",
                "Order lifecycle actions, invoice issue/void/payments, and the PDF tax-invoice download",
            ),
            (
                "Purchasing",
                "Suppliers (2), PurchaseOrders (6), SupplierInvoices (6)",
                "Order lifecycle, goods receipts (Staff may post), supplier-invoice approval and payments",
            ),
            (
                "Reporting & admin",
                "Dashboard (2), Reports (1), Company (2), AuditLogs (1)",
                "Dashboard figures and sales trend; inventory valuation export; company profile; audit trail (Administrator-only)",
            ),
            (
                "Documents",
                "Attachments (4)",
                "Upload, list, download and delete a file against any record, referenced by entity type and id; the browser client uses it on the four document detail pages",
            ),
        ],
        widths=[1.5, 3.0, 2.4],
    )
    caption(doc, "Table 9 — The 21 controllers and 83 endpoints, grouped by module.")
    para(
        doc,
        "Lifecycle actions are modelled as sub-resource POSTs rather than status fields in a PUT — for "
        "example POST /api/purchaseorders/{id}/confirm, /cancel, /receipts — so an illegal transition is "
        "a 409 from the domain, not a silently accepted update. A health endpoint at /health reports live "
        "database connectivity and is used as the deployment health check.",
    )

    # ============================================================ 12
    h1(doc, "12. The domain model")
    para(
        doc,
        "Every business entity inherits a common base: a GUID primary key, created/modified stamps with "
        "the acting user, a soft-delete marker applied through a global query filter, and a concurrency "
        "token. The entities, grouped as the Domain project groups them:",
    )
    table(
        doc,
        ["Area", "Entities", "Notes"],
        [
            (
                "Catalog",
                "Product, Category, UnitOfMeasure, Currency",
                "Product carries pricing, reorder level and its tax",
            ),
            (
                "Taxation",
                "Tax, TaxRate (+ TaxTreatment enum)",
                "Standard / ZeroRated / Exempt; Standard taxes hold effective-dated rate history",
            ),
            (
                "Organization",
                "Branch, Warehouse, CompanyProfile",
                "Warehouses belong to branches; branch scoping flows from this",
            ),
            (
                "Inventory",
                "InventoryItem, StockLayer, StockMovement (+ MovementType enum)",
                "InventoryItem is stock per product per warehouse; StockMovement is the immutable ledger; InsufficientStockException guards over-issue",
            ),
            (
                "Sales",
                "Customer, SalesOrder + SalesOrderLine, Invoice + InvoiceLine, Payment",
                "Invoice lines snapshot the VAT rate at issue; Payment records tender type",
            ),
            (
                "Purchasing",
                "Supplier, PurchaseOrder + PurchaseOrderLine, GoodsReceipt, SupplierInvoice, SupplierPayment",
                "GoodsReceipt posting creates the FIFO StockLayer at PO unit cost",
            ),
            (
                "Identity & auditing",
                "RefreshToken, LoginHistory, AuditLog",
                "Refresh tokens stored SHA-256-hashed; AuditLog rows are written by the persistence interceptor. MFA state (authenticator key, recovery codes) is held in ASP.NET Identity's own store, not a Domain entity.",
            ),
            (
                "Documents",
                "Attachment",
                "Generic file metadata keyed by (EntityType, EntityId) rather than a foreign key per module; file bytes are held by a pluggable storage adapter, not this row",
            ),
        ],
        widths=[1.4, 2.9, 2.6],
    )
    caption(doc, "Table 10 — The domain entities. State enums are listed with their aggregates in Table 5.")

    # ============================================================ 13
    h1(doc, "13. Data and persistence")
    para(
        doc,
        "Persistence is Entity Framework Core. Three database providers are supported, selected by the "
        "Database:Provider setting: SQLite for zero-infrastructure local development (the database file "
        "is created from the model on first run), SQL Server, and PostgreSQL. The PostgreSQL migrations "
        "live in their own project, ERP.Persistence.Migrations.Postgres, so the two providers' migration "
        "histories cannot contaminate each other. Domain entities carry no persistence attributes; all "
        "mapping lives in configuration classes.",
    )
    para(
        doc,
        "Three cross-cutting behaviours are applied at the persistence layer rather than trusted to each "
        "feature. An audit interceptor captures every insert, update and delete of a business entity into "
        "the audit log, with the acting user and the changed values. A global query filter hides "
        "soft-deleted rows everywhere by default. And a concurrency token on every entity turns a "
        "lost-update race into an explicit conflict instead of a silent overwrite. All access is through "
        "parameterized queries; there is no string-concatenated SQL anywhere in the solution.",
    )
    para(
        doc,
        "In the Development environment a demo data seeder populates a realistic dataset — company, "
        "branches, warehouses, catalogue, trading partners and open documents — plus the three demo "
        "accounts listed in section 15. The seeder runs only in Development.",
    )
    para(
        doc,
        "File attachments sit outside the EF Core model by design: the Attachment row is metadata, and "
        "the bytes are held by an IFileStorage port implemented today as a local-disk adapter. Storage "
        "keys are server-generated, never derived from a caller-supplied file name, so there is no "
        "path-traversal surface from an attacker-chosen name.",
    )

    # ============================================================ 14
    h1(doc, "14. Testing")
    para(
        doc,
        "The solution carries two test projects. ERP.UnitTests exercises the domain and application "
        "layers directly: document state machines, FIFO layer consumption, tax-rate selection by date, "
        "the segregation-of-duties rules, and the architecture test that fails the build if the Domain "
        "project acquires an outward dependency. ERP.IntegrationTests boots the real API in memory "
        "against a real database and drives it over HTTP, so authorization attributes, middleware, "
        "validation and persistence are all exercised together — including that a Staff token is "
        "refused where it should be and accepted for goods receipts, and that a login challenge for an "
        "MFA-enabled account is redeemable with a real TOTP code computed the same way an authenticator "
        "app would, not a test-only bypass.",
    )
    para(
        doc,
        "Across the two projects there are 97 test methods (xUnit facts and theories; theories expand "
        "to more cases at run time). Run them with `dotnet test` at the repository root.",
    )
    para(
        doc,
        "The browser client carries its own suite: Vitest with React Testing Library, 37 tests across "
        "money/date formatting, the API error-message unwrapping, and — the two places a silent "
        "regression would matter most — the full authentication state machine (silent re-auth on boot, "
        "the login/MFA-challenge/verify sequence against a mocked API, the forced-logout event) and the "
        "sign-in form itself driven the way a person would, through react-hook-form and zod validation, "
        "not by calling internal functions directly. Run them with `npm run test` in client/.",
    )
    para(
        doc,
        "A GitHub Actions workflow runs the backend suite, the frontend suite, a frontend lint pass, "
        "and both builds, on every push and pull request against main — so a regression is caught before "
        "it merges, not after someone notices it in the running application. That lint pass includes "
        "accessibility rules (oxlint's jsx-a11y plugin): every form field's label is associated with its "
        "control so a screen reader can announce what it's asking for, interactive elements are "
        "reachable by keyboard, and pages avoid disorienting focus jumps.",
    )

    # ============================================================ 15
    h1(doc, "15. Running the system")
    h2(doc, "15.1 Local development")
    para(
        doc,
        "Local development needs the .NET 9 SDK and Node 20 or later. Docker is optional — it is required "
        "only if you want SQL Server and Redis, or containerised deployment. On Windows, the run-dev.bat "
        "script at the repository root starts the API and the client in their own terminals and opens a "
        "browser. Run manually, the API starts with `dotnet run --project src/ERP.API` and listens on "
        "http://localhost:5126, creating a SQLite database from the model on first run; the client starts "
        "with `npm install` then `npm run dev` in the client directory and serves http://localhost:5173, "
        "proxying API calls so that no CORS or development-certificate configuration is needed.",
    )
    para(
        doc,
        "In the Development environment the demo data seeder creates three accounts. These credentials are "
        "seeded only in development and must never exist in a deployed system.",
    )
    table(
        doc,
        ["Email", "Password", "Role"],
        [
            ("admin@erp.local", "Admin#12345", "Administrator"),
            ("manager@erp.local", "Demo#12345", "Manager"),
            ("staff@erp.local", "Demo#12345", "Staff"),
        ],
        widths=[2.4, 2.0, 2.0],
    )
    caption(
        doc,
        "Table 11 — Development-only demo sign-ins. Walk the purchasing flow with two different accounts: "
        "segregation of duties will refuse an approval by the same person who raised the order.",
    )

    h2(doc, "15.2 Deployment")
    para(
        doc,
        "For a containerised deployment, a Dockerfile at the root builds the API and one in the client "
        "directory builds the SPA behind nginx. Secrets are supplied as environment variables and never "
        "committed: the JWT signing key, the SQL connection string, the optional bootstrap administrator, "
        "and the allowed browser origins. Interactive API documentation is available at /swagger in "
        "development, and a health endpoint at /health reports live database connectivity.",
    )
    para(
        doc,
        "A ready-made cloud deployment ships as a Render blueprint (render.yaml at the repository root). "
        "It creates three services: a managed PostgreSQL database, the API container, and the client "
        "container. The browser talks only to the client: nginx serves the SPA and proxies /api and "
        "/hubs to the API over Render's private network, so the two look like one origin — no CORS, and "
        "the API's URL is never handed to the browser. The JWT signing key is generated by Render and "
        "never seen by a human; the bootstrap administrator's email and password are deliberately "
        "prompted for on first deploy rather than generated, because the application refuses to boot "
        "with a weak administrator password rather than seed a guessable one.",
    )
    para(
        doc,
        "Error tracking is optional and off unless configured: supplying a Sentry DSN (Sentry__Dsn on "
        "the API; VITE_SENTRY_DSN, baked in at client build time, on the browser client) turns it on. "
        "Left blank, both SDKs are fully disabled and every call into them is a no-op, so this costs "
        "nothing for a deployment without a Sentry project.",
    )

    # ============================================================ 16
    h1(doc, "16. Roadmap")
    para(
        doc,
        "The following would need to be true before this system carried a real business's records. They "
        "are listed in the order a cautious operator would tackle them.",
    )
    bullets(
        doc,
        [
            "Verify the FRCS TPOS / VMS specification and implement a real IFiscalizationService adapter, including accredited invoice numbering. Until this exists, invoices remain NotSubmitted.",
            "Move the refresh token out of localStorage and into an HttpOnly, Secure, SameSite cookie, and add a policy requiring multi-factor authentication for Administrator accounts rather than leaving it opt-in.",
            "Extend the attachments screen to the remaining record types (products, customers, suppliers) beyond the four document detail pages it already covers, and add email verification of accounts. Move Hangfire from in-memory storage to a durable store so queued work survives a restart.",
            "Target real-time notifications at the users who should receive them, rather than broadcasting every event to every signed-in staff member.",
            "Decide the accounting boundary: either add a general ledger with double-entry posting from the operational documents, or define and build a clean export to the accounting package the business already uses.",
            "Implement effective-dated exchange rates and, if the business needs them, bank reconciliation against Fiji bank statement formats and the RBF payment rails.",
            "For an ephemeral or multi-instance deployment, substitute a cloud storage adapter for file attachments (local disk today) via the existing IFileStorage abstraction.",
            "Production operations: SQL Server or PostgreSQL with a tested backup and restore, secrets in a managed vault rather than environment variables, centralised log shipping alongside the optional Sentry error tracking, and a load test of the reporting queries against a realistic data volume.",
        ],
    )
    para(
        doc,
        "Nothing in this section describes work that has been started. It describes work that has not.",
        italic=True,
    )

    doc.add_paragraph()
    para(
        doc,
        "Document generated from source by docs/build_report.py on 11 July 2026.",
        size=9,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size:,} bytes)")
